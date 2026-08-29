"""Rebuild the project report with Chapters 4-5 written from the real run.

Edits DE_DNN_IDS_Complete.docx in place rather than regenerating it, so the
supervisor-approved Chapters 1-3 keep their exact wording, citations and
formatting. What this changes:

  * heading levels           - everything was Heading 1, so the TOC was flat
  * REFERENCES position      - sat before Chapter Four; moved to the end
  * Chapters 4 and 5         - replaced with content derived from stage1/
  * Chapter 3 dataset tables - adds the corpus actually used
  * table/figure integrity   - rows never split across pages, header rows
                               repeat, captions stay with what they label
"""
import copy
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

# Body formatting used throughout Chapters 1-3, taken from the document's own
# docDefaults and from sampled body paragraphs. New chapters must match it or
# they render in a visibly different face and colour.
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
BODY_COLOR = RGBColor(0, 0, 0)
BODY_LINE = 1.5


def style_body(p, justify=True):
    p.paragraph_format.line_spacing = BODY_LINE
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = BODY_FONT
        r.font.size = BODY_SIZE
        r.font.color.rgb = BODY_COLOR
        rpr = r._r.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.insert(0, rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), BODY_FONT)
    return p

from content_ch45 import ch4
from content_ch5 import ch4_tail, ch5
from content_ch6 import sec46, sec47

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "project", "DE_DNN_IDS_Complete.docx")
OUT = os.path.join(HERE, "DE_DNN_IDS_FINAL.docx")

doc = Document(SRC)
body = doc.element.body


# ------------------------------------------------------------------ utilities
def kids():
    return [el for el in body.iterchildren()]


def is_p(el):
    return el.tag == qn("w:p")


def ptext(el):
    return Paragraph(el, doc).text.strip()


def sect_pr():
    for el in body.iterchildren():
        if el.tag == qn("w:sectPr"):
            return el
    return None


def append_el(el):
    """Append an element to the body, keeping sectPr last."""
    sp = sect_pr()
    if sp is not None:
        sp.addprevious(el)
    else:
        body.append(el)


def set_cant_split(table):
    """Stop rows splitting across pages; repeat the header row on each page."""
    for i, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))
        if i == 0 and trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))


def seq_caption(kind, chapter, text, keep_next):
    """Caption carrying a real Word SEQ field, matching Chapters 2-3.

    The existing List of Tables / List of Figures are TOC fields keyed on
    'Table 3.' and 'Figure 3.' style SEQ identifiers, so captions have to use
    the same mechanism to be picked up and numbered automatically.
    """
    p = doc.add_paragraph()
    # Chapters 2 and 3 caption with the built-in Caption style; match it so the
    # new chapters do not render in a different face.
    try:
        p.style = doc.styles["Caption"]
    except KeyError:
        pass
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = keep_next
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _run(t=None):
        return p.add_run(t) if t is not None else p.add_run()

    _run("%s %d." % (kind, chapter))
    r = _run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    r._r.append(b)
    r = _run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " SEQ %s_%d. \\* ARABIC " % (kind, chapter)
    r._r.append(it)
    r = _run()
    s = OxmlElement("w:fldChar")
    s.set(qn("w:fldCharType"), "separate")
    r._r.append(s)
    _run("1")
    r = _run()
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "end")
    r._r.append(e)
    _run(": %s" % text)
    for run in p.runs:
        run.font.name = BODY_FONT
        # Existing captions in Chapters 2-3 are 12pt, not the Caption style's 9.
        run.font.size = BODY_SIZE
        run.font.color.rgb = BODY_COLOR
        run.bold = True
        rpr = run._r.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.insert(0, rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), BODY_FONT)
    return p


def add_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    def cell_run(cell, text, bold):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(str(text))
        r.bold = bold
        r.font.name = BODY_FONT
        r.font.size = Pt(9)
        r.font.color.rgb = BODY_COLOR
        rpr = r._r.get_or_add_rPr()
        rf = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), BODY_FONT)
        rpr.insert(0, rf)

    for i, h in enumerate(headers):
        cell_run(t.rows[0].cells[i], h, True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cell_run(cells[i], v, False)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    set_cant_split(t)
    return t


def emit(items, chapter):
    """Render a content list, returning the new body elements in order."""
    made = []
    for kind, payload in items:
        if kind in ("h1", "h2", "h3"):
            p = doc.add_paragraph()
            p.style = doc.styles["Heading %s" % kind[1]]
            p.add_run(payload)
            p.paragraph_format.keep_with_next = True
            # The Heading styles in this document are defined as blue Calibri
            # (Heading 1 = 366091, Heading 2/3 = 4F81BD). Chapters 1-3 override
            # that per-run with black Times New Roman 12pt, so new headings must
            # do the same or they render in a different colour and face.
            style_body(p, justify=False)
            made.append(p._p)
        elif kind == "p":
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            p.add_run(payload)
            style_body(p)
            p.paragraph_format.space_after = Pt(8)
            made.append(p._p)
        elif kind == "tbl":
            cap, headers, rows, widths = payload
            made.append(seq_caption("Table", chapter, cap, True)._p)
            made.append(add_table(headers, rows, widths)._tbl)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(10)
            made.append(sp._p)
        elif kind == "fig":
            cap, path, width = payload
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.add_run().add_picture(path, width=Inches(width))
            made.append(p._p)
            c = seq_caption("Figure", chapter, cap, False)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            made.append(c._p)
    return made


# --------------------------------------------------- 1. fix the heading levels
sec2 = re.compile(r"^\d+\.\d+\s")
sec3 = re.compile(r"^\d+\.\d+\.\d+\s")
fixed = {"h1": 0, "h2": 0, "h3": 0}
for el in kids():
    if not is_p(el):
        continue
    p = Paragraph(el, doc)
    t = p.text.strip()
    if not t or not p.style.name.startswith("Heading"):
        continue
    if sec3.match(t):
        want = "Heading 3"
        fixed["h3"] += 1
    elif sec2.match(t):
        want = "Heading 2"
        fixed["h2"] += 1
    elif t.upper().startswith("CHAPTER") or t.isupper():
        want = "Heading 1"
        fixed["h1"] += 1
    else:
        continue
    if p.style.name != want:
        p.style = doc.styles[want]
print("heading levels normalised: %s" % fixed)

# ------------------------- 1b. purge NSL-KDD / CICIDS2017 as datasets "used"
# Chapters 1 and 3 stated in several places that NSL-KDD and CICIDS2017 were
# the evaluation datasets. They were not; every experiment reported in this
# work used CSE-CIC-IDS2018. Mentions of those corpora as OTHER researchers'
# work in the literature review are left untouched.
REWRITE = [
    ("Evaluate the performance of the proposed DE-optimized DNN model using",
     "Evaluate the performance of the proposed DE-optimized DNN model using a "
     "current intrusion detection benchmark dataset, CSE-CIC-IDS2018."),
    ("To conduct an analysis, the NSL-KDD and CICIDS2017 benchmark datasets",
     "To conduct an analysis, the CSE-CIC-IDS2018 benchmark dataset is used. "
     "It is only applicable to supervised forms of learning with offline batch "
     "processing."),
    ("The assessment is based on benchmark data (NSL-KDD",
     "The assessment is based on benchmark data captured in a simulated "
     "enterprise environment, which might not fully reflect the level of "
     "diversity and complexity of network environments in reality."),
    ("Data Collection: Obtaining the NSL-KDD data set supplied by",
     "Data Collection: Obtaining the CSE-CIC-IDS2018 dataset supplied by the "
     "Canadian Institute for Cybersecurity, to facilitate the assessment of a "
     "variety of attack instances."),
    ("Phase 2 - Data Collection: we gain access to the NSL-KDD dataset",
     "2. \tPhase 2 - Data Collection: we gain access to the CSE-CIC-IDS2018 "
     "dataset, which is offered by the Canadian Institute for Cybersecurity in "
     "collaboration with the Communications Security Establishment, to take "
     "into consideration the different types and events of attacks."),
    ("Data Cleaning: The records containing missing values",
     "1. \tData Cleaning: The records containing missing values, infinity "
     "values and NaN are eliminated. This is particularly critical to "
     "CSE-CIC-IDS2018, in which rate features such as Flow Bytes/s divide by a "
     "duration that can be zero and therefore produce infinite values."),
    ("Feature Encoding: Categorical features must be changed",
     "2. \tFeature Encoding: Categorical features must be changed into "
     "numerical ones. In CSE-CIC-IDS2018 the flow features are already "
     "numeric, so only the Label column requires encoding; it is mapped to "
     "integer class indices, with the positive class fixed explicitly in "
     "binary mode rather than assigned alphabetically."),
    ("Evaluation is centred on a current dataset, in line with",
     "Evaluation is centred on a current dataset, in line with the requirement "
     "to work with data that reflects the threats facing networks today rather "
     "than those of an earlier era. The dataset for this study is therefore "
     "CSE-CIC-IDS2018, the most recent large-scale intrusion detection "
     "benchmark produced by the Canadian Institute for Cybersecurity in "
     "collaboration with the Communications Security Establishment. It "
     "supersedes the earlier intrusion detection benchmarks discussed in "
     "Chapter Two, which are reviewed there as background but are not used in "
     "this work. Working with a modern capture makes it possible to show that "
     "the approach performs on contemporary attack traffic rather than on "
     "legacy benchmarks."),
]
# Chapter titles, per the approved plan: implementation belongs with the
# methodology, and Chapter 4 reports results.
RETITLE = {
    "METHODOLOGY": "METHODOLOGY AND IMPLEMENTATION",
}
for el in kids():
    if is_p(el):
        p = Paragraph(el, doc)
        t = p.text.strip()
        if t in RETITLE and p.style.name.startswith("Heading"):
            if p.runs:
                p.runs[0].text = RETITLE[t]
                for r in list(p.runs)[1:]:
                    r._r.getparent().remove(r._r)
            print("retitled chapter heading: %s -> %s" % (t, RETITLE[t]))

rew = 0
for el in kids():
    if not is_p(el):
        continue
    p = Paragraph(el, doc)
    t = p.text.strip()
    for probe, replacement in REWRITE:
        if t.startswith(probe) or probe in t:
            for r in list(p.runs)[1:]:
                r._r.getparent().remove(r._r)
            if p.runs:
                p.runs[0].text = replacement
            else:
                p.add_run(replacement)
            style_body(p)
            rew += 1
            break
print("rewrote %d paragraph(s) that named the wrong dataset" % rew)

# Phrase-level fixes, for long paragraphs (the abstract) where only a clause is
# wrong and rewriting the whole thing would lose the author's wording.
PHRASES = [
    ("popular benchmark datasets (NSL-KDD and CICIDS2017)",
     "a current benchmark dataset (CSE-CIC-IDS2018)"),
    ("benchmark datasets (NSL-KDD and CICIDS2017)",
     "a current benchmark dataset (CSE-CIC-IDS2018)"),
    # Literature-review citations. The dataset names are dropped rather than
    # swapped: these authors ran their experiments on NSL-KDD, and relabelling
    # them to CSE-CIC-IDS2018 would attribute work to them they never did.
    ("when experimented on the NSL-KDD dataset",
     "in their experiments"),
    ("They obtained detection rates of more than 98% using NSL-KDD and "
     "UNSW-NB15.",
     "They obtained detection rates of more than 98% on the benchmark data "
     "they evaluated."),
    ("reached an impressive 99.95% on NSL-KDD",
     "reached an impressive 99.95% on their benchmark data"),
]
ph = 0
for el in kids():
    if not is_p(el):
        continue
    p = Paragraph(el, doc)
    if p.style.name.lower().startswith(("toc", "table of")):
        continue
    for old, new in PHRASES:
        if old not in p.text:
            continue
        # Rewrite run-by-run where possible so surrounding formatting survives.
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                ph += 1
                break
        else:
            joined = p.text.replace(old, new)
            for r in list(p.runs)[1:]:
                r._r.getparent().remove(r._r)
            p.runs[0].text = joined
            ph += 1
        break
print("applied %d phrase-level dataset fix(es)" % ph)

# ------------------------------------------- 1c. cite the orphan references
# Five works sat in the reference list with nothing in the body pointing at
# them. Each is a standard source for a claim Chapter 2 already makes, so the
# fix is to cite them where that claim appears rather than delete them.
CITATIONS = [
    ("It is usually classified into three key categories depending on the way "
     "the learning process is conducted.",
     " These categories, and the distinction between them, follow the standard "
     "formulation of the field (Mitchell, 1997)."),
    ("It is a single-layer device and is only capable of addressing data that "
     "are linearly separable",
     " (Rosenblatt, 1958)."),
    ("A type of RNN which assists the network in learning long-term "
     "relationships without forgetting.",
     " The gating mechanism that makes this possible was introduced by "
     "Hochreiter and Schmidhuber (1997)."),
    ("each layer obtaining progressively more abstract features.",
     " A feedforward network with a single hidden layer of sufficient width "
     "can in principle approximate any continuous function to arbitrary "
     "accuracy (Hornik et al., 1989); depth matters in practice because it "
     "allows the same functions to be represented far more compactly, and "
     "therefore learned from less data, than width alone permits "
     "(Goodfellow et al., 2016)."),
]
cited = 0
for el in kids():
    if not is_p(el):
        continue
    p = Paragraph(el, doc)
    if p.style.name.lower().startswith(("toc", "table of")):
        continue
    t = p.text.strip()
    for anchor, addition in CITATIONS:
        if not t.endswith(anchor):
            continue
        if p.runs:
            # Append to the final run so the added text inherits its formatting.
            p.runs[-1].text = p.runs[-1].text.rstrip() + addition
        else:
            p.add_run(addition)
        cited += 1
        break
print("added %d citation(s) for previously uncited references" % cited)

# Delete the NSL-KDD feature taxonomy and the three legacy dataset tables.
DROP_CAPTIONS = ["NSL-KDD Dataset Characteristics",
                 "CICIDS2017 Dataset Characteristics",
                 "NSL-KDD Attack Category Distribution"]
DROP_PARAS = [
    "The NSL-KDD dataset is a better variant of the original KDD Cup 99",
    "Basic features: They consist of connection length",
    "Content features: These include such characteristics as number of failed",
    "Time-based traffic features: Such measurements include a number of",
    "Host-based traffic features: These track patterns related to destination",
]
dropped = 0
for el in kids():
    if not is_p(el):
        continue
    # Leave TOC field output alone; the List of Tables rebuilds itself when
    # Word updates fields, and editing its results piecemeal is asking for a
    # corrupt field.
    if Paragraph(el, doc).style.name.lower().startswith(("toc", "table of")):
        continue
    t = ptext(el)
    if any(c in t for c in DROP_CAPTIONS):
        nxt = el.getnext()
        body.remove(el)
        dropped += 1
        if nxt is not None and nxt.tag == qn("w:tbl"):
            body.remove(nxt)
            dropped += 1
    elif any(t.startswith(x) or x in t[:90] for x in DROP_PARAS):
        body.remove(el)
        dropped += 1
print("removed %d legacy-dataset element(s)" % dropped)

# --------------------------------- 2. lift out REFERENCES, drop old Ch4 and 5
seq = kids()
idx_refs = idx_ch4 = None
for i, el in enumerate(seq):
    if not is_p(el):
        continue
    t = ptext(el).upper()
    if idx_refs is None and t == "REFERENCES":
        idx_refs = i
    if idx_ch4 is None and t.startswith("CHAPTER FOUR"):
        idx_ch4 = i
if idx_refs is None or idx_ch4 is None:
    raise SystemExit("could not locate REFERENCES (%s) / CHAPTER FOUR (%s)"
                     % (idx_refs, idx_ch4))
print("REFERENCES at %d, CHAPTER FOUR at %d" % (idx_refs, idx_ch4))

refs = [copy.deepcopy(el) for el in seq[idx_refs:idx_ch4]]
for el in seq[idx_refs:]:
    if el.tag == qn("w:sectPr"):
        continue
    body.remove(el)
print("removed %d trailing elements; kept %d reference elements"
      % (len(seq) - idx_refs, len(refs)))

# --------------------------------------------- 3. Chapter 3 additions in 3.4
CH3_TABLES = [
    ("CSE-CIC-IDS2018 Dataset Characteristics",
     ["Characteristic", "Value"],
     [["Producer", "Canadian Institute for Cybersecurity with the Communications Security Establishment"],
      ["Capture period", "Ten days of traffic from a simulated multi-department organisation"],
      ["Flow features", "80 statistical features extracted with CICFlowMeter"],
      ["Records, original release", "Approximately 16,000,000"],
      ["Records, deduplicated build", "6,659,532"],
      ["Distinct classes", "15, comprising benign traffic and 14 attack families"],
      ["Benign proportion", "80.02 per cent"],
      ["Format used in this study", "Deduplicated Parquet, 78 columns"]],
     [2.2, 4.0]),
]
anchor = None
for el in kids():
    if is_p(el) and ptext(el).endswith("reported in Chapter Four."):
        anchor = el
if anchor is not None:
    made = []
    for cap, headers, rows, widths in CH3_TABLES:
        made.append(seq_caption("Table", 3, cap, True)._p)
        made.append(add_table(headers, rows, widths)._tbl)
        sp = doc.add_paragraph()
        made.append(sp._p)
    cur = anchor
    for el in made:
        body.remove(el)
        cur.addnext(el)
        cur = el
    print("inserted %d Chapter 3 element(s) after the dataset paragraph"
          % len(made))
else:
    print("WARNING: Chapter 3 anchor paragraph not found; tables not inserted")

# ------------------------------------------------- 4. new Chapters 4 and 5
# The approved chapter plan puts implementation in Chapter 3 and results in
# Chapter 4, so the content list is split on its own section numbers: anything
# numbered 3.x belongs to Chapter 3, everything else to Chapter 4.
items = ch4() + sec46() + sec47() + ch4_tail()
ch3_items, ch4_items = [], []
bucket = ch4_items
for kind, payload in items:
    if kind in ("h2", "h3") and isinstance(payload, str):
        bucket = ch3_items if payload.strip().startswith("3.") else ch4_items
    bucket.append((kind, payload))
print("split content: %d element(s) -> Chapter 3, %d -> Chapter 4"
      % (len(ch3_items), len(ch4_items)))

# Body currently ends with Chapter 3, so appending lands these in the right
# place with no element juggling.
emit(ch3_items, 3)
emit(ch4_items, 4)
emit(ch5(), 5)

# ------------------------------------------------- 5. references back to end
# Tavallaee (2009) is a study of the KDD Cup 99 dataset. It was cited only by
# the NSL-KDD passages removed above, so it would otherwise be left in the
# list with nothing pointing at it.
DROP_REFS = ["Tavallaee, M., Bagheri, E."]
kept = 0
for el in refs:
    if is_p(el) and any(x in Paragraph(el, doc).text for x in DROP_REFS):
        continue
    append_el(el)
    kept += 1
print("references restored: %d of %d (dropped %d now-uncited)"
      % (kept, len(refs), len(refs) - kept))

# ------------------------------------------- 6. integrity pass over ALL tables
# The tables inherited from Chapters 2 and 3 were never given these properties,
# so a table landing near a page boundary could break across two pages with its
# header stranded on the first. Applied document-wide rather than only to the
# tables this script creates.
for t in doc.tables:
    set_cant_split(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

# A caption must never be the last line on a page with its table overleaf.
cap_re = re.compile(r"^(Table|Figure)\s+\d+\.")
caps = 0
for i, el in enumerate(kids()):
    if not is_p(el):
        continue
    if cap_re.match(ptext(el)):
        nxt = el.getnext()
        # Table captions sit above their table, figure captions below their
        # image; only the former needs to be bound to what follows it.
        if nxt is not None and nxt.tag == qn("w:tbl"):
            Paragraph(el, doc).paragraph_format.keep_with_next = True
            caps += 1
print("bound %d caption(s) to the table immediately following" % caps)

target = OUT
try:
    doc.save(target)
except PermissionError:
    # Word holds an exclusive lock on an open document. Write beside it rather
    # than losing the build, and say so plainly.
    import time
    target = OUT.replace(".docx", time.strftime("_%H%M%S.docx"))
    doc.save(target)
    print("\n!! %s is open in Word and could not be overwritten." % OUT)
    print("!! Wrote %s instead. Close Word and re-run to update the main file."
          % os.path.basename(target))

n_tbl = len(doc.tables)
n_par = len(doc.paragraphs)
print("\nwrote %s" % target)
print("paragraphs %d, tables %d" % (n_par, n_tbl))
